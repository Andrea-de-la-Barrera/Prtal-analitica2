import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from pathlib import Path
from datetime import timedelta

# ✅ Word report
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ✅ Toggle global (botón + abrir/cerrar sidebar)
from ui.sidebar_toggle import sidebar_toggle

# ==========================================
# CONFIG
# ==========================================
st.set_page_config(
    page_title="Analítica Descriptiva",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Bootstrap icons
st.markdown(
    '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.1/font/bootstrap-icons.css">',
    unsafe_allow_html=True
)

# ==========================================
# Cargar CSS (global + módulo)
# ==========================================
base_dir = Path(__file__).resolve().parents[1]
global_css = base_dir / "styles" / "global.css"
module_css = base_dir / "styles" / "descriptiva.css"

css_text = ""
if global_css.exists():
    css_text += global_css.read_text(encoding="utf-8") + "\n"
if module_css.exists():
    css_text += module_css.read_text(encoding="utf-8") + "\n"

# ✅ Extra CSS: letras blancas para captions / texto auxiliar
css_text += """
/* captions en blanco */
.stCaption, .stCaption *{ color: rgba(226,232,240,0.88) !important; }
/* markdown (texto) ligeramente más claro */
.stMarkdown, .stMarkdown *{ color: rgba(255,255,255,0.95); }
"""

if css_text.strip():
    st.markdown(f"<style>{css_text}</style>", unsafe_allow_html=True)
else:
    st.warning("No se encontró CSS (global.css / descriptiva.css).")

# ✅ Toggle después del CSS
sidebar_toggle()

# ==========================================
# Helpers
# ==========================================
MESES_ES = ["ene", "feb", "mar", "abr", "may", "jun", "jul", "ago", "sep", "oct", "nov", "dic"]

def money(x) -> str:
    try:
        return f"${float(x):,.2f}"
    except Exception:
        return str(x)

def style_dark_matplotlib(ax):
    """Estilo dark para matplotlib sin pelear con tu UI."""
    ax.set_facecolor((0, 0, 0, 0))
    ax.figure.set_facecolor((0, 0, 0, 0))
    ax.tick_params(colors="white")
    ax.xaxis.label.set_color("white")
    ax.yaxis.label.set_color("white")
    ax.title.set_color("white")
    ax.grid(True, alpha=0.18)
    for spine in ax.spines.values():
        spine.set_alpha(0.25)
        spine.set_color("white")

def format_date_es(d):
    """Devuelve fecha tipo: 03 mar 2025"""
    if pd.isna(d):
        return ""
    d = pd.to_datetime(d).date()
    return f"{d.day:02d} {MESES_ES[d.month-1]} {d.year}"

def make_period_labels(df: pd.DataFrame, nivel: str) -> pd.DataFrame:
    """
    Crea columna 'periodo' para mensual o semanal + etiqueta legible.
    """
    out = df.copy()
    if nivel == "Mensual":
        out["periodo"] = out["fecha"].dt.to_period("M").astype(str)  # 2025-08
        y = out["fecha"].dt.year
        m = out["fecha"].dt.month
        out["periodo_etq"] = m.map(lambda mm: MESES_ES[mm-1]) + " " + y.astype(str)
        return out

    iso = out["fecha"].dt.isocalendar()
    out["iso_year"] = iso["year"].astype(int)
    out["iso_week"] = iso["week"].astype(int)

    start = out["fecha"] - pd.to_timedelta(out["fecha"].dt.weekday, unit="D")
    end = start + pd.to_timedelta(6, unit="D")

    out["semana_inicio"] = start.dt.date
    out["semana_fin"] = end.dt.date

    out["periodo"] = out["iso_year"].astype(str) + "-W" + out["iso_week"].astype(str).str.zfill(2)

    out["periodo_etq"] = out.apply(
        lambda r: f"Semana {int(r['iso_week'])} ({format_date_es(r['semana_inicio'])[:-5]} – {format_date_es(r['semana_fin'])[:-5]})",
        axis=1
    )
    return out

def build_insights(resumen_periodo: pd.DataFrame, resumen_categoria: pd.Series, unidad: str) -> dict:
    """
    Interpretaciones dinámicas SIN markdown (**).
    resumen_periodo: columnas -> periodo, periodo_etq, gasto_total
    unidad: "mes" o "semana"
    """
    insights = {}

    m = resumen_periodo.copy()
    m["gasto_total"] = pd.to_numeric(m["gasto_total"], errors="coerce").fillna(0)

    n = int(len(m))
    unidad_txt = "mes" if unidad == "mes" else "semana"
    unidad_pl = "meses" if unidad == "mes" else "semanas"

    first = float(m["gasto_total"].iloc[0])
    last = float(m["gasto_total"].iloc[-1])
    cambio_pct = ((last - first) / first * 100) if first != 0 else 0.0

    x = np.arange(n)
    y = m["gasto_total"].values.astype(float)
    slope_pct = 0.0
    if n >= 2:
        slope = np.polyfit(x, y, 1)[0]
        y_mean = float(np.mean(y)) if float(np.mean(y)) != 0 else 1.0
        slope_pct = (float(slope) / y_mean) * 100

    if slope_pct > 2:
        trend_txt = "tendencia al alza"
    elif slope_pct < -2:
        trend_txt = "tendencia a la baja"
    else:
        trend_txt = "tendencia estable"

    vol = (float(np.std(y)) / (float(np.mean(y)) if float(np.mean(y)) != 0 else 1.0)) * 100
    if vol >= 25:
        vol_txt = "alta variabilidad (gasto irregular)"
    elif vol >= 12:
        vol_txt = "variabilidad moderada"
    else:
        vol_txt = "gasto relativamente constante"

    i_max = int(m["gasto_total"].values.argmax())
    i_min = int(m["gasto_total"].values.argmin())

    periodo_max = str(m["periodo_etq"].iloc[i_max])
    val_max = float(m["gasto_total"].iloc[i_max])

    periodo_min = str(m["periodo_etq"].iloc[i_min])
    val_min = float(m["gasto_total"].iloc[i_min])

    jump_txt = ""
    if n >= 3:
        dif = pd.Series(y).pct_change().replace([np.inf, -np.inf], np.nan).dropna()
        if not dif.empty:
            j = int(dif.abs().values.argmax())
            idx_to = int(dif.index[j])
            idx_from = idx_to - 1

            salto = float(dif.iloc[j] * 100)
            if abs(salto) >= 20:
                p_from = str(m["periodo_etq"].iloc[idx_from])
                p_to = str(m["periodo_etq"].iloc[idx_to])
                jump_txt = f"Se detecta un cambio brusco entre {p_from} → {p_to} ({salto:+.1f}%)."

    insights["periodo"] = (
        f"En el periodo analizado ({n} {unidad_pl}), el gasto presenta {trend_txt}. "
        f"Del primer al último {unidad_txt} hay un cambio de {cambio_pct:+.1f}%. "
        f"Además, se observa {vol_txt}. "
        f"El mayor gasto fue en {periodo_max} ({money(val_max)}), y el menor en {periodo_min} ({money(val_min)}). "
        f"{jump_txt}".strip()
    )

    cat = resumen_categoria.copy()
    total = float(cat.sum()) if float(cat.sum()) != 0 else 1.0

    top1 = str(cat.index[0]) if len(cat) else "N/A"
    top1_val = float(cat.iloc[0]) if len(cat) else 0.0
    top1_pct = (top1_val / total) * 100

    top3_val = float(cat.head(3).sum()) if len(cat) >= 3 else float(cat.sum())
    top3_pct = (top3_val / total) * 100

    if top1_pct >= 40:
        conc_txt = "El gasto está muy concentrado en una sola categoría."
    elif top1_pct >= 25:
        conc_txt = "El gasto está concentrado en la categoría principal."
    else:
        conc_txt = "El gasto está distribuido entre varias categorías."

    insights["categoria"] = (
        f"La categoría con mayor gasto es {top1} con {money(top1_val)} ({top1_pct:.1f}% del total). "
        f"Las 3 principales categorías representan {top3_pct:.1f}% del gasto. "
        f"{conc_txt}"
    )

    return insights

def format_table_for_display(df: pd.DataFrame, money_cols=None) -> pd.DataFrame:
    money_cols = money_cols or []
    out = df.copy()

    for c in money_cols:
        if c in out.columns:
            out[c] = out[c].apply(money)

    return out

# ============== WORD HELPERS ==============
def fig_to_png_bytes(fig) -> BytesIO:
    img = BytesIO()
    fig.savefig(img, format="png", dpi=200, bbox_inches="tight")
    img.seek(0)
    return img

def add_df_table_to_doc_full(doc: Document, df: pd.DataFrame, title: str, money_cols=None):
    """
    Inserta la tabla COMPLETA en Word.
    OJO: si df tiene miles de filas, el Word será pesado.
    """
    money_cols = set(money_cols or [])
    doc.add_heading(title, level=2)

    if df.empty:
        doc.add_paragraph("Sin datos para mostrar.")
        return

    # tabla completa
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    # filas
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if col in money_cols:
                cells[j].text = money(val)
            else:
                cells[j].text = str(val)

def build_word_report_full(
    df_filtrado: pd.DataFrame,
    resumen_periodo_show: pd.DataFrame,
    resumen_categoria: pd.Series,
    insights: dict,
    nivel: str,
    start_date,
    end_date,
    fig_evolucion=None,
    fig_categorias=None
) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Reporte de Analítica Descriptiva", level=0)
    p = doc.add_paragraph("Proyecto: Portal de Analítica (ANALITICA)\n")
    p.add_run(f"Nivel de análisis: {nivel}\n")
    p.add_run(f"Rango de fechas: {start_date} a {end_date}\n")
    doc.add_paragraph("")

    # KPIs
    doc.add_heading("KPIs", level=1)
    gasto_total = float(df_filtrado["monto"].sum())
    promedio = float(df_filtrado["monto"].mean())
    maximo = float(df_filtrado["monto"].max())
    registros = int(len(df_filtrado))

    for line in [
        f"Gasto total: {money(gasto_total)}",
        f"Promedio: {money(promedio)}",
        f"Máximo: {money(maximo)}",
        f"Registros: {registros}",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph("")

    # Insights
    doc.add_heading("Interpretación automática", level=1)
    doc.add_paragraph(insights.get("periodo", ""))
    doc.add_paragraph(insights.get("categoria", ""))
    doc.add_paragraph("")

    # Gráficas
    doc.add_heading("Visualizaciones", level=1)

    if fig_evolucion is not None:
        doc.add_paragraph("Evolución del gasto")
        img1 = fig_to_png_bytes(fig_evolucion)
        doc.add_picture(img1, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    if fig_categorias is not None:
        doc.add_paragraph("Gasto total por categoría")
        img2 = fig_to_png_bytes(fig_categorias)
        doc.add_picture(img2, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    # Tablas
    add_df_table_to_doc_full(
        doc,
        resumen_periodo_show,
        title="Resumen del periodo",
        money_cols=["gasto_total", "gasto_promedio"]
    )

    doc.add_heading("Ranking por categoría", level=1)
    cat_df = resumen_categoria.reset_index()
    cat_df.columns = ["categoria", "monto_total"]
    add_df_table_to_doc_full(
        doc,
        cat_df,
        title="Ranking completo por categoría",
        money_cols=["monto_total"]
    )

    # Detalle completo (todas las filas filtradas)
    add_df_table_to_doc_full(
        doc,
        df_filtrado,
        title="Detalle completo de movimientos filtrados",
        money_cols=["monto"]
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ==========================================
# HERO
# ==========================================
st.markdown(
    """
    <div class="hero-wrap">
        <div class="pill"><span class="pill-dot"></span> EQ. BTS (DESCRIPTIVA)</div>
        <div class="hero-title">Analítica <span class="accent">Descriptiva</span></div>
        <div class="hero-sub">
            Sube tu CSV, filtra por fechas y categorías, visualiza KPIs y revisa insights automáticos.
        </div>
    </div>
    <div class="sep"></div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# PANEL REQUISITOS
# ==========================================
st.markdown(
    """
    <div class="panel">
        <div class="panel-row">
            <div class="panel-icon">
                <i class="bi bi-info-circle"></i>
            </div>
            <div>
                <div class="panel-title">Requisitos del archivo</div>
                <small>
                    <b>Columnas requeridas:</b>
                    <code>fecha</code>, <code>categoria</code>, <code>monto</code>
                    &nbsp;&nbsp;|&nbsp;&nbsp;
                    <b>Opcionales:</b>
                    <code>metodo_pago</code>, <code>nota</code>
                </small>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# CARGA CSV
# ==========================================
file = st.file_uploader("Sube tu archivo CSV", type=["csv"])
if not file:
    st.info("Sube un CSV para comenzar.")
    st.stop()

df = pd.read_csv(file)
df.columns = [c.strip().lower() for c in df.columns]

required = {"fecha", "categoria", "monto"}
if not required.issubset(set(df.columns)):
    st.error(f"El CSV debe contener al menos estas columnas: {', '.join(sorted(required))}.")
    st.stop()

# Limpieza básica
df["fecha"] = pd.to_datetime(df["fecha"], errors="coerce")
df["monto"] = pd.to_numeric(df["monto"], errors="coerce")
df["categoria"] = df["categoria"].astype(str).str.strip()

df = df.dropna(subset=["fecha", "categoria", "monto"])

if df.empty:
    st.warning("El archivo no tiene datos válidos después de limpiar (fecha/categoria/monto).")
    st.stop()

# ==========================================
# FILTROS
# ==========================================
st.sidebar.markdown("### Filtros")

rango = st.sidebar.date_input(
    "Rango de fechas",
    (df["fecha"].min().date(), df["fecha"].max().date())
)
start, end = rango

df_f = df[(df["fecha"].dt.date >= start) & (df["fecha"].dt.date <= end)].copy()

cats = sorted(df_f["categoria"].dropna().unique())
sel_cats = st.sidebar.multiselect("Categorías", cats, default=cats)
df_f = df_f[df_f["categoria"].isin(sel_cats)].copy()

if df_f.empty:
    st.warning("Con los filtros actuales no hay datos para mostrar. Ajusta el rango o categorías.")
    st.stop()

# ==========================================
# KPIs
# ==========================================
c1, c2, c3, c4 = st.columns(4)
c1.metric("Gasto total", money(df_f["monto"].sum()))
c2.metric("Promedio", money(df_f["monto"].mean()))
c3.metric("Máximo", money(df_f["monto"].max()))
c4.metric("Registros", int(len(df_f)))

st.divider()

# ==========================================
# NIVEL DE ANÁLISIS (Mensual / Semanal)
# ==========================================
nivel = st.radio(
    "Nivel de análisis",
    ["Mensual", "Semanal"],
    horizontal=True,
    index=0
)

df_p = make_period_labels(df_f, nivel=nivel)

# ==========================================
# AGREGACIONES (Periodo + Categoría)
# ==========================================
resumen_periodo = (
    df_p.groupby(["periodo", "periodo_etq"], as_index=False)
        .agg(
            gasto_total=("monto", "sum"),
            gasto_promedio=("monto", "mean"),
            movimientos=("monto", "count")
        )
        .sort_values("periodo")
)

resumen_categoria = (
    df_f.groupby("categoria")["monto"].sum().sort_values(ascending=False)
)

unidad = "mes" if nivel == "Mensual" else "semana"
insights = build_insights(resumen_periodo, resumen_categoria, unidad=unidad)

# ==========================================
# GRÁFICAS
# ==========================================
colA, colB = st.columns(2)

with colA:
    st.markdown(f"## Evolución del gasto {nivel.lower()}")

    fig, ax = plt.subplots(figsize=(7.4, 4.2))
    ax.plot(resumen_periodo["periodo_etq"], resumen_periodo["gasto_total"], marker="o", linewidth=2)
    ax.set_xlabel("Periodo")
    ax.set_ylabel("Monto ($)")
    ax.set_title(f"Gasto total por {nivel.lower()}")
    plt.xticks(rotation=35, ha="right")

    style_dark_matplotlib(ax)
    plt.tight_layout()
    st.pyplot(fig, use_container_width=True)

with colB:
    st.markdown("## Gasto total por categoría")

    fig2, ax2 = plt.subplots(figsize=(7.4, 4.2))
    ax2.bar(resumen_categoria.index.astype(str), resumen_categoria.values.astype(float))
    ax2.set_xlabel("Categoría")
    ax2.set_ylabel("Monto ($)")
    ax2.set_title("Ranking de gasto por categoría")
    plt.xticks(rotation=35, ha="right")

    style_dark_matplotlib(ax2)
    plt.tight_layout()
    st.pyplot(fig2, use_container_width=True)

# ==========================================
# INSIGHTS
# ==========================================
st.markdown(
    f"""
    <div class="insights-grid">
      <div class="insight-card">
        <div class="insight-title">Interpretación automática — {nivel}</div>
        <p class="insight-text">{insights["periodo"]}</p>
      </div>
      <div class="insight-card">
        <div class="insight-title">Interpretación automática — Categoría</div>
        <p class="insight-text">{insights["categoria"]}</p>
      </div>
    </div>
    """,
    unsafe_allow_html=True
)

st.divider()

# ==========================================
# TABLAS (con scroll)
# ==========================================
st.markdown("## Resumen del periodo (tabla)")

resumen_show = pd.DataFrame({
    "periodo": resumen_periodo["periodo_etq"].astype(str),
    "gasto_total": resumen_periodo["gasto_total"],
    "gasto_promedio": resumen_periodo["gasto_promedio"],
    "movimientos": resumen_periodo["movimientos"],
})

st.dataframe(
    format_table_for_display(resumen_show, money_cols=["gasto_total", "gasto_promedio"]),
    use_container_width=True,
    height=260
)

st.markdown("## Detalle de movimientos filtrados")
detalle = df_f.copy()
detalle["fecha"] = detalle["fecha"].dt.strftime("%Y-%m-%d")

st.dataframe(
    format_table_for_display(detalle, money_cols=["monto"]),
    use_container_width=True,
    height=420
)

# ==========================================
# DESCARGA (EXCEL CON FORMATO + GRÁFICAS)
# ==========================================
from io import BytesIO
from openpyxl.styles import Font, PatternFill
from openpyxl.chart import LineChart, BarChart, Reference

excel_buffer = BytesIO()

with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
    # ----------------------------------
    # HOJA 1: RESUMEN
    # ----------------------------------
    resumen_excel = resumen_show.copy()

    resumen_excel.to_excel(
        writer,
        sheet_name="Resumen_periodo",
        index=False
    )

    # ----------------------------------
    # HOJA 2: DETALLE
    # ----------------------------------
    detalle_excel = df_f.copy()
    # 👉 fecha como fecha REAL de Excel
    detalle_excel["fecha"] = pd.to_datetime(detalle_excel["fecha"])

    detalle_excel.to_excel(
        writer,
        sheet_name="Detalle_movimientos",
        index=False
    )

    wb = writer.book

    # ==================================
    # ESTILOS COMUNES
    # ==================================
    header_fill = PatternFill(
        start_color="1F4FD8", end_color="1F4FD8", fill_type="solid"
    )  # azul rey
    header_font = Font(color="FFFFFF", bold=True)

    # ----------------------------------
    # FORMATO RESUMEN
    # ----------------------------------
    ws_r = writer.sheets["Resumen_periodo"]
    ws_r.freeze_panes = "A2"

    for cell in ws_r[1]:
        cell.fill = header_fill
        cell.font = header_font

    for col in ["B", "C"]:
        for cell in ws_r[col][1:]:
            cell.number_format = '"$"#,##0.00'

    ws_r.column_dimensions["A"].width = 24
    ws_r.column_dimensions["B"].width = 18
    ws_r.column_dimensions["C"].width = 20
    ws_r.column_dimensions["D"].width = 15

    # ----------------------------------
    # FORMATO DETALLE
    # ----------------------------------
    ws_d = writer.sheets["Detalle_movimientos"]
    ws_d.freeze_panes = "A2"

    for cell in ws_d[1]:
        cell.fill = header_fill
        cell.font = header_font

    ws_d.column_dimensions["A"].width = 14  # fecha
    ws_d.column_dimensions["B"].width = 22  # categoria
    ws_d.column_dimensions["C"].width = 16  # monto
    ws_d.column_dimensions["D"].width = 22
    ws_d.column_dimensions["E"].width = 30

    for cell in ws_d["A"][1:]:
        cell.number_format = "yyyy-mm-dd"

    for cell in ws_d["C"][1:]:
        cell.number_format = '"$"#,##0.00'

    # ==================================
    # GRÁFICA 1: EVOLUCIÓN DEL GASTO
    # ==================================
    chart_line = LineChart()
    chart_line.title = "Evolución del gasto"
    chart_line.y_axis.title = "Monto ($)"
    chart_line.x_axis.title = "Periodo"

    data = Reference(ws_r, min_col=2, min_row=1, max_row=ws_r.max_row)
    cats = Reference(ws_r, min_col=1, min_row=2, max_row=ws_r.max_row)

    chart_line.add_data(data, titles_from_data=True)
    chart_line.set_categories(cats)
    chart_line.height = 10
    chart_line.width = 20

    ws_r.add_chart(chart_line, "F2")

    # ==================================
    # GRÁFICA 2: GASTO POR CATEGORÍA
    # ==================================
    cat_df = resumen_categoria.reset_index()
    cat_df.columns = ["categoria", "monto_total"]

    cat_df.to_excel(
        writer,
        sheet_name="Categorias",
        index=False
    )

    ws_c = writer.sheets["Categorias"]

    for cell in ws_c[1]:
        cell.fill = header_fill
        cell.font = header_font

    ws_c.column_dimensions["A"].width = 22
    ws_c.column_dimensions["B"].width = 18

    for cell in ws_c["B"][1:]:
        cell.number_format = '"$"#,##0.00'

    chart_bar = BarChart()
    chart_bar.title = "Gasto total por categoría"
    chart_bar.y_axis.title = "Monto ($)"
    chart_bar.x_axis.title = "Categoría"

    data = Reference(ws_c, min_col=2, min_row=1, max_row=ws_c.max_row)
    cats = Reference(ws_c, min_col=1, min_row=2, max_row=ws_c.max_row)

    chart_bar.add_data(data, titles_from_data=True)
    chart_bar.set_categories(cats)
    chart_bar.height = 10
    chart_bar.width = 20

    ws_c.add_chart(chart_bar, "D2")

excel_buffer.seek(0)

st.download_button(
    "⬇ Descargar Excel (con gráficas)",
    data=excel_buffer.getvalue(),
    file_name="reporte_descriptivo_filtrado.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)


# ==========================================
# DESCARGA
# ==========================================


st.caption("Generar Word con detalle completo puede tardar si hay muchos registros.")

# ✅ Generación bajo demanda (para no congelar en cada rerun)
if "word_bytes" not in st.session_state:
    st.session_state.word_bytes = None

if st.button("📄 Generar Word (detalle completo)"):
    # Generar Word con el df filtrado completo (no 'detalle' string)
    word_df = df_f.copy()
    word_df["fecha"] = word_df["fecha"].dt.strftime("%Y-%m-%d")

    st.session_state.word_bytes = build_word_report_full(
        df_filtrado=word_df,
        resumen_periodo_show=resumen_show,
        resumen_categoria=resumen_categoria,
        insights=insights,
        nivel=nivel,
        start_date=start,
        end_date=end,
        fig_evolucion=fig,
        fig_categorias=fig2
    )

if st.session_state.word_bytes:
    st.download_button(
        "⬇ Descargar Reporte Word (.docx)",
        data=st.session_state.word_bytes,
        file_name="reporte_descriptivo_completo.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.caption("© 2025 Portal de Analítica | Módulo Descriptivo")
