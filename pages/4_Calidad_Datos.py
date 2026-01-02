import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from pathlib import Path
from io import BytesIO
from datetime import datetime

# ‚úÖ Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ‚úÖ Toggle global (bot√≥n + abrir/cerrar sidebar)
from ui.sidebar_toggle import sidebar_toggle


# ==========================================
# CONFIG
# ==========================================
st.set_page_config(
    page_title="Calidad de Datos",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Bootstrap icons
st.markdown(
    '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.1/font/bootstrap-icons.css">',
    unsafe_allow_html=True
)

# ==========================================
# Cargar CSS (global + m√≥dulo)
# ==========================================
base_dir = Path(__file__).resolve().parents[1]  # .../portal_analitica
global_css = base_dir / "styles" / "global.css"
module_css = base_dir / "styles" / "calidad_datos.css"

css_text = ""
if global_css.exists():
    css_text += global_css.read_text(encoding="utf-8") + "\n"
if module_css.exists():
    css_text += module_css.read_text(encoding="utf-8") + "\n"

if css_text.strip():
    st.markdown(f"<style>{css_text}</style>", unsafe_allow_html=True)
else:
    st.warning("No se encontr√≥ CSS (global.css / calidad_datos.css).")

# ‚úÖ Toggle despu√©s del CSS
sidebar_toggle()

# ==========================================
# Helpers
# ==========================================
def style_dark_matplotlib(ax):
    ax.set_facecolor((0, 0, 0, 0))
    ax.figure.set_facecolor((0, 0, 0, 0))
    ax.tick_params(colors="white")
    ax.xaxis.label.set_color("white")
    ax.yaxis.label.set_color("white")
    ax.title.set_color("white")
    ax.grid(True, alpha=0.18)
    for spine in ax.spines.values():
        spine.set_alpha(0.25)
        spine.set_color("white")

def score_level(score: float):
    if score >= 90:
        return "Excelente"
    if score >= 75:
        return "Buena"
    if score >= 60:
        return "Regular"
    return "Riesgosa"

def clamp(v, lo, hi):
    return max(lo, min(hi, v))

def safe_pct(n, total):
    return 0.0 if total <= 0 else (n / total) * 100

def compute_outliers_iqr(series: pd.Series, factor: float):
    s = pd.to_numeric(series, errors="coerce").dropna()
    if len(s) < 5:
        return 0, 0.0
    q1 = float(s.quantile(0.25))
    q3 = float(s.quantile(0.75))
    iqr = q3 - q1
    if iqr == 0:
        return 0, 0.0
    lower = q1 - factor * iqr
    upper = q3 + factor * iqr
    out_count = int(((s < lower) | (s > upper)).sum())
    out_pct = safe_pct(out_count, len(series))
    return out_count, round(out_pct, 2)

def build_quality_explain(score, nivel, dup_count, dup_pct, avg_null_pct, avg_out_pct):
    bullets = []

    if avg_null_pct > 10:
        bullets.append("<b>Nulos altos:</b> revisa columnas clave y define estrategia (imputar/eliminar/validar).")
    elif avg_null_pct > 0:
        bullets.append("<b>Nulos detectados:</b> est√°n en nivel manejable; revisa si afectan columnas clave.")
    else:
        bullets.append("‚úÖ <b>Sin nulos relevantes:</b> buena se√±al para modelado.")

    if dup_count > 0:
        bullets.append(f"<b>Duplicados:</b> hay {dup_count} filas duplicadas ({dup_pct:.2f}%). Elim√≠nalos para evitar sesgos.")
    else:
        bullets.append("‚úÖ <b>Sin duplicados:</b> evita conteos inflados y sesgos.")

    if avg_out_pct > 5:
        bullets.append("<b>Outliers:</b> hay valores extremos; valida si son errores o casos reales (pueden mover modelos).")
    elif avg_out_pct > 0:
        bullets.append("<b>Outliers ligeros:</b> revisa si son esperables en tu contexto.")
    else:
        bullets.append("‚úÖ <b>Sin outliers relevantes:</b> datos m√°s estables para modelos.")

    if nivel in ["Excelente", "Buena"]:
        next_step = "Tu dataset est√° bastante listo. Limpia lo m√≠nimo (si aplica) y pasa a Predictiva / Miner√≠a."
    elif nivel == "Regular":
        next_step = "Haz una limpieza b√°sica (nulos/duplicados/outliers) y vuelve a correr este m√≥dulo."
    else:
        next_step = "Primero corrige calidad (nulos/duplicados/tipos/outliers). Si no, Predictiva/Miner√≠a dar√° resultados raros."

    exec_line = (
        f"Resumen: Score {score:.2f}/100 | "
        f"Duplicados: {dup_count} ({dup_pct:.2f}%) | "
        f"Nulos prom.: {avg_null_pct:.2f}% | "
        f"Outliers prom.: {avg_out_pct:.2f}%."
    )
    return exec_line, bullets, next_step


# ==========================
# Word helpers
# ==========================
def fig_to_png_bytes(fig) -> bytes:
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()

def add_df_to_doc(doc: Document, df: pd.DataFrame, title: str, max_rows: int = 60):
    doc.add_heading(title, level=2)

    if df is None or df.empty:
        p = doc.add_paragraph("Sin datos para mostrar.")
        p.runs[0].italic = True
        return

    show = df.copy()
    if len(show) > max_rows:
        show = show.head(max_rows)
        doc.add_paragraph(f"(Mostrando primeras {max_rows} filas)")

    # Convertir todo a string para evitar errores en docx
    show = show.fillna("").astype(str)

    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(show.columns):
        hdr_cells[j].text = str(col)

    for _, row in show.iterrows():
        row_cells = table.add_row().cells
        for j, val in enumerate(row.values):
            row_cells[j].text = str(val)

def make_quality_docx_bytes(
    df: pd.DataFrame,
    preview_rows: int,
    dtypes_df: pd.DataFrame,
    nulos_df: pd.DataFrame,
    outliers_df: pd.DataFrame,
    dup_count: int,
    dup_pct: float,
    avg_null_pct: float,
    avg_out_pct: float,
    score: float,
    nivel: str,
    exec_line: str,
    bullets_plain: list,
    next_step: str,
    advanced: bool,
    params: dict,
    fig_nulls_png: bytes | None,
    fig_out_png: bytes | None,
):
    doc = Document()

    # Title
    doc.add_heading("Reporte ‚Äî Calidad de Datos", level=0)
    meta = doc.add_paragraph()
    meta.add_run("Generado: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))

    doc.add_paragraph(" ")

    # Resumen ejecutivo
    doc.add_heading("Resumen ejecutivo", level=1)
    p = doc.add_paragraph()
    p.add_run("Score: ").bold = True
    p.add_run(f"{score:.2f}/100")
    p = doc.add_paragraph()
    p.add_run("Nivel: ").bold = True
    p.add_run(nivel)

    doc.add_paragraph(exec_line)

    # Hallazgos (bullets)
    doc.add_heading("Hallazgos principales", level=1)
    for b in bullets_plain:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Siguiente paso recomendado", level=1)
    doc.add_paragraph(next_step)

    # Configuraci√≥n usada
    doc.add_heading("Configuraci√≥n usada", level=1)
    doc.add_paragraph(f"Modo avanzado: {'S√≠' if advanced else 'No'}")
    for k, v in params.items():
        doc.add_paragraph(f"{k}: {v}")

    # Preview
    add_df_to_doc(doc, df.head(int(preview_rows)), "Vista previa del dataset")

    # Tipos
    add_df_to_doc(doc, dtypes_df, "Tipos de datos detectados")

    # Nulos
    add_df_to_doc(doc, nulos_df, "Valores nulos por columna", max_rows=80)

    # Duplicados
    doc.add_heading("Filas duplicadas", level=2)
    doc.add_paragraph(f"Duplicados detectados: {dup_count} ({dup_pct:.2f}%)")

    # Outliers
    add_df_to_doc(doc, outliers_df, "Outliers (IQR) en columnas num√©ricas", max_rows=80)

    # Gr√°ficas
    doc.add_heading("Gr√°ficas", level=1)

    if fig_nulls_png:
        doc.add_paragraph("Nulos (Top 12 columnas):")
        doc.add_picture(BytesIO(fig_nulls_png), width=Inches(6.3))

    if fig_out_png:
        doc.add_paragraph("Outliers (Top 12 num√©ricas):")
        doc.add_picture(BytesIO(fig_out_png), width=Inches(6.3))

    doc.add_paragraph(" ")
    foot = doc.add_paragraph("¬© 2025 Portal de Anal√≠tica | M√≥dulo Calidad de Datos")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ==========================================
# HERO
# ==========================================
st.markdown(
    """
    <div class="hero-wrap">
        <div class="pill"><span class="pill-dot"></span> EQ. BTS (Calidad)</div>
        <div class="hero-title">Calidad de <span class="accent">Datos</span></div>
        <div class="hero-sub">
            Eval√∫a tu dataset antes de anal√≠tica o miner√≠a:
            <b>nulos</b>, <b>duplicados</b>, <b>tipos</b> y <b>outliers</b>.
        </div>
    </div>
    <div class="sep"></div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# SIDEBAR: gu√≠a + configuraci√≥n
# ==========================================
st.sidebar.markdown(
    """
    <div class="sidebar-help">
        <h4>¬øC√≥mo usar este m√≥dulo?</h4>
        <ul>
            <li>Sube tu CSV (cualquier estructura).</li>
            <li>Revisa tipos, nulos, duplicados y outliers.</li>
            <li>Activa <b>Modo avanzado</b> si deseas ajustar sensibilidad.</li>
            <li>Descarga el reporte antes de modelar.</li>
        </ul>
        <div class="sidebar-tip">
            Tip: si Predictiva/Miner√≠a da m√©tricas raras, casi siempre la causa est√° aqu√≠.
        </div>
    </div>
    <div class="sidebar-sep"></div>
    <h4 class="sidebar-title">Configuraci√≥n</h4>
    """,
    unsafe_allow_html=True
)

preview_rows = st.sidebar.number_input(
    "Filas a previsualizar",
    min_value=5,
    max_value=200,
    value=30,
    step=5
)

advanced = st.sidebar.toggle("Modo avanzado (sensibilidad del score)", value=False)

# Defaults ‚Äúseguros‚Äù
w_null = 1.2
w_dup = 0.7
w_out = 1.9
iqr_factor = 1.7
min_n_numeric = 16

if advanced:
    st.sidebar.markdown("### Sensibilidad (score)")
    w_null = st.sidebar.slider("Peso Nulos", 0.0, 3.0, float(w_null), 0.05)
    w_dup = st.sidebar.slider("Peso Duplicados", 0.0, 3.0, float(w_dup), 0.05)
    w_out = st.sidebar.slider("Peso Outliers", 0.0, 3.0, float(w_out), 0.05)

    st.sidebar.markdown("### Outliers (IQR)")
    iqr_factor = st.sidebar.slider("Factor IQR", 1.0, 3.5, float(iqr_factor), 0.05)
    min_n_numeric = st.sidebar.number_input("M√≠n. datos por columna num√©rica", 5, 200, int(min_n_numeric), 1)

# ==========================================
# PANEL REQUISITOS
# ==========================================
st.markdown(
    """
    <div class="panel">
        <div style="display:flex; gap:12px;">
            <div class="panel-icon"><i class="bi bi-shield-check"></i></div>
            <div>
                <div class="panel-title">¬øQu√© revisa este m√≥dulo?</div>
                <small>
                    <b>Entrada:</b> CSV con cualquier estructura &nbsp;|&nbsp;
                    <b>Salida:</b> diagn√≥stico + score + recomendaciones + reporte descargable
                </small>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# CARGA CSV
# ==========================================
file = st.file_uploader("Sube tu archivo CSV (cualquier estructura)", type=["csv"])
if not file:
    st.info("Sube un CSV para comenzar.")
    st.stop()

df = None
read_errors = []

for kwargs in [
    dict(sep=",", encoding="utf-8"),
    dict(sep=";", encoding="utf-8"),
    dict(sep=",", encoding="latin-1"),
    dict(sep=";", encoding="latin-1"),
]:
    try:
        file.seek(0)
        df = pd.read_csv(file, **kwargs)
        break
    except Exception as e:
        read_errors.append(str(e))

if df is None:
    st.error("No se pudo leer el CSV con separadores comunes (, ;) y encodings (utf-8/latin-1).")
    st.code("\n\n".join(read_errors[:3]))
    st.stop()

if df.empty:
    st.warning("El archivo est√° vac√≠o o no contiene registros.")
    st.stop()

df.columns = [str(c).strip() for c in df.columns]

# ==========================================
# VISTA PREVIA
# ==========================================
st.markdown('<div class="section-title">Vista previa del dataset</div>', unsafe_allow_html=True)
st.dataframe(df.head(int(preview_rows)), use_container_width=True)
st.divider()

# ==========================================
# 1) TIPOS DE DATOS
# ==========================================
st.markdown('<div class="section-title">1) Tipos de datos detectados</div>', unsafe_allow_html=True)

dtypes_df = pd.DataFrame({
    "columna": df.columns,
    "tipo_detectado": [str(df[c].dtype) for c in df.columns],
    "valores_unicos": [int(df[c].nunique(dropna=True)) for c in df.columns]
})
st.dataframe(dtypes_df, use_container_width=True)

# ==========================================
# 2) NULOS
# ==========================================
st.markdown('<div class="section-title">2) Valores nulos por columna</div>', unsafe_allow_html=True)
nulos = df.isna().sum()
porc_nulos = (nulos / len(df)) * 100

nulos_df = pd.DataFrame({
    "columna": nulos.index.astype(str),
    "nulos": nulos.values.astype(int),
    "%_nulos": np.round(porc_nulos.values, 2),
}).sort_values("%_nulos", ascending=False)

st.dataframe(nulos_df, use_container_width=True)

# ==========================================
# 3) DUPLICADOS
# ==========================================
st.markdown('<div class="section-title">3) Filas duplicadas</div>', unsafe_allow_html=True)
dup_count = int(df.duplicated().sum())
dup_pct = safe_pct(dup_count, len(df))
st.write(f"**Filas duplicadas detectadas:** {dup_count} ({dup_pct:.2f}%)")

# ==========================================
# 4) OUTLIERS (IQR) EN NUM√âRICAS
# ==========================================
st.markdown('<div class="section-title">4) Outliers (IQR) en columnas num√©ricas</div>', unsafe_allow_html=True)

num_cols = df.select_dtypes(include=[np.number]).columns.tolist()

out_rows = []
if not num_cols:
    st.info("No se detectaron columnas num√©ricas para an√°lisis de outliers.")
    outliers_df = pd.DataFrame(columns=["columna", "outliers_detectados", "%_outliers"])
else:
    for col in num_cols:
        series = df[col]
        if pd.to_numeric(series, errors="coerce").dropna().shape[0] < int(min_n_numeric):
            out_rows.append([str(col), 0, 0.0])
            continue
        out_count, out_pct = compute_outliers_iqr(series, float(iqr_factor))
        out_rows.append([str(col), out_count, out_pct])

    outliers_df = (
        pd.DataFrame(out_rows, columns=["columna", "outliers_detectados", "%_outliers"])
        .sort_values("%_outliers", ascending=False)
        .reset_index(drop=True)
    )

st.dataframe(outliers_df, use_container_width=True)
st.divider()

# ==========================================
# 5) SCORE CALIDAD (ESTIMADO)
# ==========================================
st.markdown('<div class="section-title">5) Score de Calidad (estimado)</div>', unsafe_allow_html=True)

avg_null_pct = float(nulos_df["%_nulos"].mean()) if len(nulos_df) else 0.0
avg_out_pct = float(outliers_df["%_outliers"].mean()) if len(outliers_df) else 0.0

penalty = (avg_null_pct * float(w_null)) + (dup_pct * float(w_dup)) + (avg_out_pct * float(w_out))
score = clamp(100.0 - penalty, 0.0, 100.0)
nivel = score_level(score)

c1, c2, c3, c4 = st.columns(4)
c1.metric("Score (0‚Äì100)", f"{score:.2f}")
c2.metric("Nivel", nivel)
c3.metric("Nulos promedio", f"{avg_null_pct:.2f}%")
c4.metric("Outliers promedio", f"{avg_out_pct:.2f}%")

# ==========================================
# VISUALIZACI√ìN (calidad)
# ==========================================
st.markdown('<div class="section-title">Visualizaci√≥n (calidad)</div>', unsafe_allow_html=True)

g1, g2 = st.columns(2)

fig_nulls_png = None
fig_out_png = None

with g1:
    st.markdown("### % Nulos (Top 12 columnas)")
    top_nulls = nulos_df.head(12).copy()
    fig, ax = plt.subplots(figsize=(7.4, 4.2))
    ax.bar(top_nulls["columna"].astype(str), top_nulls["%_nulos"].astype(float))
    ax.set_ylabel("% nulos")
    ax.set_title("Columnas con m√°s nulos")
    plt.xticks(rotation=35, ha="right")
    style_dark_matplotlib(ax)
    plt.tight_layout()
    st.pyplot(fig, use_container_width=True)
    fig_nulls_png = fig_to_png_bytes(fig)

with g2:
    st.markdown("### % Outliers (Top 12 num√©ricas)")
    if len(outliers_df) == 0:
        st.info("No hay columnas num√©ricas.")
    else:
        top_out = outliers_df.head(12).copy()
        fig2, ax2 = plt.subplots(figsize=(7.4, 4.2))
        ax2.bar(top_out["columna"].astype(str), top_out["%_outliers"].astype(float))
        ax2.set_ylabel("% outliers")
        ax2.set_title("Columnas con m√°s outliers")
        plt.xticks(rotation=35, ha="right")
        style_dark_matplotlib(ax2)
        plt.tight_layout()
        st.pyplot(fig2, use_container_width=True)
        fig_out_png = fig_to_png_bytes(fig2)

# ==========================================
# INTERPRETACI√ìN PRO (CARD)
# ==========================================
exec_line, bullets, next_step = build_quality_explain(
    score=score,
    nivel=nivel,
    dup_count=dup_count,
    dup_pct=dup_pct,
    avg_null_pct=avg_null_pct,
    avg_out_pct=avg_out_pct
)

badge_class = "badge-ok" if nivel in ["Excelente", "Buena"] else ("badge-warn" if nivel == "Regular" else "badge-bad")
bullets_html = "".join([f"<li>{b}</li>" for b in bullets])

card_html = f"""
<div class="interpret-card">
  <div class="interpret-head">
    <div class="interpret-title">Interpretaci√≥n autom√°tica</div>
    <span class="badge {badge_class}">{nivel}</span>
  </div>

  <div class="interpret-exec">{exec_line}</div>

  <div class="interpret-text">
    Este m√≥dulo sirve para detectar problemas t√≠picos de calidad que luego provocan
    m√©tricas raras, modelos inestables o conclusiones equivocadas.
  </div>

  <ul class="interpret-bullets">
    {bullets_html}
  </ul>

  <div class="interpret-tip">
    <span class="tip-label">Siguiente paso:</span> {next_step}
  </div>

  <div class="interpret-note">
    Recomendado <b>antes</b> de Predictiva / Miner√≠a para evitar sesgos,
    errores por tipos incorrectos o resultados inestables por datos extremos.
  </div>
</div>
"""
st.markdown(card_html, unsafe_allow_html=True)

# ==========================================
# DESCARGAS
# ==========================================
st.divider()
st.markdown('<div class="section-title">Descargar reporte</div>', unsafe_allow_html=True)

report_cols = pd.DataFrame({
    "columna": dtypes_df["columna"].astype(str),
    "tipo_detectado": dtypes_df["tipo_detectado"].astype(str),
    "valores_unicos": dtypes_df["valores_unicos"].astype(int),
}).merge(
    nulos_df[["columna", "nulos", "%_nulos"]],
    on="columna",
    how="left"
)

if len(outliers_df):
    report_cols = report_cols.merge(outliers_df, on="columna", how="left")
else:
    report_cols["outliers_detectados"] = 0
    report_cols["%_outliers"] = 0.0

report_cols["outliers_detectados"] = report_cols["outliers_detectados"].fillna(0).astype(int)
report_cols["%_outliers"] = report_cols["%_outliers"].fillna(0.0).astype(float)

# Resumen global (constantes)
report_cols["filas_total"] = int(len(df))
report_cols["duplicados_total"] = int(dup_count)
report_cols["duplicados_pct"] = float(round(dup_pct, 2))
report_cols["nulos_promedio_pct"] = float(round(avg_null_pct, 2))
report_cols["outliers_promedio_pct"] = float(round(avg_out_pct, 2))
report_cols["score_calidad"] = float(round(score, 2))
report_cols["nivel"] = nivel

# ==========================================
# DESCARGA EXCEL (CALIDAD DE DATOS + GR√ÅFICAS)
# ==========================================
from openpyxl.styles import Font, PatternFill
from openpyxl.chart import BarChart, Reference
from io import BytesIO

excel_buffer = BytesIO()

with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
    # ----------------------------------
    # HOJA 1: REPORTE POR COLUMNA
    # ----------------------------------
    report_cols_excel = report_cols.copy()
    report_cols_excel.to_excel(
        writer,
        sheet_name="Calidad_por_columna",
        index=False
    )

    # ----------------------------------
    # HOJA 2: RESUMEN GENERAL
    # ----------------------------------
    resumen_excel = pd.DataFrame({
        "m√©trica": [
            "Filas totales",
            "Duplicados",
            "% Duplicados",
            "% Nulos promedio",
            "% Outliers promedio",
            "Score de calidad",
            "Nivel"
        ],
        "valor": [
            int(len(df)),
            int(dup_count),
            round(dup_pct, 2),
            round(avg_null_pct, 2),
            round(avg_out_pct, 2),
            round(score, 2),
            nivel
        ]
    })

    resumen_excel.to_excel(
        writer,
        sheet_name="Resumen_calidad",
        index=False
    )

    wb = writer.book

    # ==================================
    # ESTILOS COMUNES
    # ==================================
    header_fill = PatternFill(
        start_color="1F4FD8", end_color="1F4FD8", fill_type="solid"
    )  # azul rey
    header_font = Font(color="FFFFFF", bold=True)

    # ----------------------------------
    # FORMATO HOJA: CALIDAD POR COLUMNA
    # ----------------------------------
    ws_c = writer.sheets["Calidad_por_columna"]
    ws_c.freeze_panes = "A2"

    for cell in ws_c[1]:
        cell.fill = header_fill
        cell.font = header_font

    ws_c.column_dimensions["A"].width = 26
    ws_c.column_dimensions["B"].width = 18
    ws_c.column_dimensions["C"].width = 18
    ws_c.column_dimensions["D"].width = 14
    ws_c.column_dimensions["E"].width = 14
    ws_c.column_dimensions["F"].width = 20

    for col in ["E", "H", "J"]:
        if col in ws_c:
            for cell in ws_c[col][1:]:
                cell.number_format = '0.00"%"'

    # ----------------------------------
    # FORMATO HOJA: RESUMEN
    # ----------------------------------
    ws_r = writer.sheets["Resumen_calidad"]
    ws_r.freeze_panes = "A2"

    for cell in ws_r[1]:
        cell.fill = header_fill
        cell.font = header_font

    ws_r.column_dimensions["A"].width = 26
    ws_r.column_dimensions["B"].width = 22

    # ==================================
    # GR√ÅFICA 1: % NULOS (TOP 10)
    # ==================================
    top_nulls_excel = nulos_df.head(10).copy()
    top_nulls_excel.to_excel(
        writer,
        sheet_name="Top_nulos",
        index=False
    )

    ws_n = writer.sheets["Top_nulos"]
    for cell in ws_n[1]:
        cell.fill = header_fill
        cell.font = header_font

    chart_nulls = BarChart()
    chart_nulls.title = "% de nulos por columna (Top 10)"
    chart_nulls.y_axis.title = "% nulos"
    chart_nulls.x_axis.title = "Columna"

    data = Reference(
        ws_n,
        min_col=3,
        min_row=1,
        max_row=ws_n.max_row
    )
    cats = Reference(
        ws_n,
        min_col=1,
        min_row=2,
        max_row=ws_n.max_row
    )

    chart_nulls.add_data(data, titles_from_data=True)
    chart_nulls.set_categories(cats)
    chart_nulls.height = 11
    chart_nulls.width = 22

    ws_n.add_chart(chart_nulls, "E2")

    # ==================================
    # GR√ÅFICA 2: % OUTLIERS (TOP 10)
    # ==================================
    if len(outliers_df):
        top_out_excel = outliers_df.head(10).copy()
        top_out_excel.to_excel(
            writer,
            sheet_name="Top_outliers",
            index=False
        )

        ws_o = writer.sheets["Top_outliers"]
        for cell in ws_o[1]:
            cell.fill = header_fill
            cell.font = header_font

        chart_out = BarChart()
        chart_out.title = "% de outliers por columna (Top 10)"
        chart_out.y_axis.title = "% outliers"
        chart_out.x_axis.title = "Columna"

        data = Reference(
            ws_o,
            min_col=3,
            min_row=1,
            max_row=ws_o.max_row
        )
        cats = Reference(
            ws_o,
            min_col=1,
            min_row=2,
            max_row=ws_o.max_row
        )

        chart_out.add_data(data, titles_from_data=True)
        chart_out.set_categories(cats)
        chart_out.height = 11
        chart_out.width = 22

        ws_o.add_chart(chart_out, "E2")

excel_buffer.seek(0)

st.download_button(
    "‚¨á Descargar Excel (Calidad de Datos)",
    data=excel_buffer.getvalue(),
    file_name="reporte_calidad_datos.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)


# ==========================
# ‚úÖ NUEVO: WORD (TODO EL REPORTE)
# ==========================
params_used = {
    "Filas previsualizadas": int(preview_rows),
    "Peso Nulos (w_null)": float(w_null),
    "Peso Duplicados (w_dup)": float(w_dup),
    "Peso Outliers (w_out)": float(w_out),
    "Factor IQR": float(iqr_factor),
    "M√≠n. datos por columna num√©rica": int(min_n_numeric),
}

bullets_plain = [
    b.replace("<b>", "").replace("</b>", "").replace("‚úÖ", "‚Ä¢")
    for b in bullets
]

# Card visual (si ya tienes CSS, se ver√° pro)
st.markdown(
    """
    <div class="word-card">
      <div class="word-title">Reporte Word (completo)</div>
      <div class="word-sub">
        Incluye: vista previa, tipos, nulos, duplicados, outliers, score, interpretaci√≥n y gr√°ficas.
      </div>
    </div>
    """,
    unsafe_allow_html=True
)

if st.button("üìù Generar Word (reporte completo)"):
    docx_bytes = make_quality_docx_bytes(
        df=df,
        preview_rows=int(preview_rows),
        dtypes_df=dtypes_df,
        nulos_df=nulos_df,
        outliers_df=outliers_df,
        dup_count=int(dup_count),
        dup_pct=float(dup_pct),
        avg_null_pct=float(avg_null_pct),
        avg_out_pct=float(avg_out_pct),
        score=float(score),
        nivel=str(nivel),
        exec_line=str(exec_line),
        bullets_plain=bullets_plain,
        next_step=str(next_step),
        advanced=bool(advanced),
        params=params_used,
        fig_nulls_png=fig_nulls_png,
        fig_out_png=fig_out_png,
    )

    st.download_button(
        "‚¨á Descargar Word (Calidad de Datos)",
        data=docx_bytes,
        file_name="reporte_calidad_datos.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

st.caption("¬© 2025 Portal de Anal√≠tica | M√≥dulo Calidad de Datos")
