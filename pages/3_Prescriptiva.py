import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from textwrap import dedent
from pathlib import Path
from io import BytesIO

# Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ✅ Toggle global (botón + abrir/cerrar sidebar)
from ui.sidebar_toggle import sidebar_toggle


# ==========================================
# CONFIG
# ==========================================
st.set_page_config(
    page_title="Analítica Prescriptiva",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Bootstrap icons
st.markdown(
    '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.1/font/bootstrap-icons.css">',
    unsafe_allow_html=True
)

# ==========================================
# Cargar CSS (global + módulo)
# ==========================================
base_dir = Path(__file__).resolve().parents[1]
global_css = base_dir / "styles" / "global.css"
module_css = base_dir / "styles" / "prescriptiva.css"

css_text = ""
if global_css.exists():
    css_text += global_css.read_text(encoding="utf-8") + "\n"
if module_css.exists():
    css_text += module_css.read_text(encoding="utf-8") + "\n"

if css_text.strip():
    st.markdown(f"<style>{css_text}</style>", unsafe_allow_html=True)
else:
    st.warning("No se encontró CSS (global.css / prescriptiva.css).")

# ✅ Toggle después del CSS
sidebar_toggle()


# ==========================================
# Helpers
# ==========================================
def money(x):
    try:
        return f"${float(x):,.2f}"
    except Exception:
        return str(x)

def style_dark_matplotlib(ax):
    ax.set_facecolor((0, 0, 0, 0))
    ax.figure.set_facecolor((0, 0, 0, 0))
    ax.tick_params(colors="white")
    ax.xaxis.label.set_color("white")
    ax.yaxis.label.set_color("white")
    ax.title.set_color("white")
    ax.grid(True, alpha=0.18)
    for spine in ax.spines.values():
        spine.set_alpha(0.25)
        spine.set_color("white")

def build_prescriptive_explain(total_actual, presupuesto, exceso, df_plan, scenario_name):
    # df_plan: categoria, gasto_mensual, porcentaje_total, reducir_sugerido, objetivo_recomendado, prioridad, accion
    top_cat = df_plan.sort_values("gasto_mensual", ascending=False).iloc[0]
    top_red = df_plan.sort_values("reducir_sugerido", ascending=False).iloc[0]

    top_cat_name = str(top_cat["categoria"])
    top_cat_amt  = float(top_cat["gasto_mensual"])
    top_cat_pct  = float(top_cat["porcentaje_total"]) * 100

    top_red_name = str(top_red["categoria"])
    top_red_amt  = float(top_red["reducir_sugerido"])

    if exceso <= 0:
        status_title = "Dentro del presupuesto"
        status_text = (
            f"Tu gasto total estimado es {money(total_actual)} y tu presupuesto objetivo es {money(presupuesto)}. "
            "No se requiere recorte para cumplir el objetivo."
        )
        bullets = [
            f"La categoría con mayor peso es {top_cat_name}: {money(top_cat_amt)} ({top_cat_pct:.1f}% del total).",
            "Recomendación: mantén límites por categoría y monitorea mes a mes.",
            f"Escenario seleccionado: {scenario_name}. (Aquí no aplica recorte porque no hay exceso.)"
        ]
        action_hint = "Si quieres optimizar, baja ligeramente categorías no esenciales para generar margen."
    else:
        status_title = "Ajuste requerido"
        status_text = (
            f"Tu gasto total estimado es {money(total_actual)} y tu presupuesto objetivo es {money(presupuesto)}. "
            f"Para cumplirlo, se requiere un ajuste aproximado de {money(exceso)}."
        )
        bullets = [
            f"La categoría con mayor peso es {top_cat_name}: {money(top_cat_amt)} ({top_cat_pct:.1f}% del total).",
            f"El recorte principal sugerido se concentra en {top_red_name}: {money(top_red_amt)}.",
            f"Escenario seleccionado: {scenario_name}. Este escenario define qué tan fuerte y amplio se distribuye el recorte."
        ]
        action_hint = (
            "Tip práctico: empieza por las 3 categorías con prioridad ALTA. "
            "Aplica el recorte sugerido y revisa el resultado por 2–4 semanas."
        )

    executive = (
        f"Resumen: {status_title}. "
        f"Total: {money(total_actual)} | Objetivo: {money(presupuesto)} | Ajuste: {money(exceso)}."
    )
    return status_title, status_text, bullets, action_hint, executive


# =========================
# WORD HELPERS
# =========================
def fig_to_png_bytes(fig) -> BytesIO:
    img = BytesIO()
    fig.savefig(img, format="png", dpi=200, bbox_inches="tight")
    img.seek(0)
    return img

def add_df_table_to_doc(doc: Document, df: pd.DataFrame, title: str, max_rows: int = 200):
    doc.add_heading(title, level=2)
    if df is None or df.empty:
        doc.add_paragraph("Sin datos para mostrar.")
        return

    df2 = df.copy()
    if len(df2) > max_rows:
        df2 = df2.head(max_rows)

    table = doc.add_table(rows=1, cols=len(df2.columns))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for j, col in enumerate(df2.columns):
        hdr[j].text = str(col)

    for _, row in df2.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df2.columns):
            cells[j].text = str(row[col])

    if len(df) > max_rows:
        doc.add_paragraph(f"Nota: se muestran solo {max_rows} filas (de {len(df)}).")

def build_prescriptive_word_report(
    # Metadatos/estado
    registros_cargados, registros_validos, neg_detect, neg_removed, remove_negative,
    escenario, meses_label,
    presupuesto, total_actual, exceso,
    # Tablas
    df_resumen, df_show, df_out,
    # Interpretación
    status_title, status_text, bullets, action_hint, executive,
    # Figuras
    fig_comp=None, fig_ahorro=None
) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Reporte de Analítica Prescriptiva", level=0)
    doc.add_paragraph("Proyecto: Portal de Analítica (ANALITICA)")
    doc.add_paragraph(f"Escenario: {escenario}")
    doc.add_paragraph(f"Rango de meses: {meses_label}")
    doc.add_paragraph("")

    doc.add_heading("Estado del archivo", level=1)
    doc.add_paragraph(f"Registros cargados: {registros_cargados}", style="List Bullet")
    doc.add_paragraph(f"Registros válidos: {registros_validos}", style="List Bullet")
    doc.add_paragraph(f"Negativos detectados: {neg_detect}", style="List Bullet")
    doc.add_paragraph(
        f"Quitar negativos: {'Sí' if remove_negative else 'No'} | Negativos removidos: {neg_removed if remove_negative else 0}",
        style="List Bullet"
    )
    doc.add_paragraph("")

    doc.add_heading("KPIs", level=1)
    doc.add_paragraph(f"Presupuesto objetivo: {money(presupuesto)}", style="List Bullet")
    doc.add_paragraph(f"Total actual: {money(total_actual)}", style="List Bullet")
    doc.add_paragraph(f"Ajuste requerido: {money(exceso)}", style="List Bullet")
    doc.add_paragraph("")

    add_df_table_to_doc(doc, df_resumen, "Gasto actual por categoría")
    doc.add_paragraph("")

    add_df_table_to_doc(doc, df_show, "Plan prescriptivo (vista para usuario)")
    doc.add_paragraph("")

    add_df_table_to_doc(doc, df_out, "Plan prescriptivo (valores numéricos)")
    doc.add_paragraph("")

    doc.add_heading("Interpretación automática", level=1)
    doc.add_paragraph(f"Estado: {status_title}")
    doc.add_paragraph(executive)
    doc.add_paragraph(status_text)
    doc.add_paragraph("Puntos clave:")
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph("")
    doc.add_paragraph(f"Siguiente paso: {action_hint}")
    doc.add_paragraph("")

    doc.add_heading("Visualizaciones", level=1)

    if fig_comp is not None:
        doc.add_paragraph("Gasto actual vs objetivo recomendado (Top 10)")
        doc.add_picture(fig_to_png_bytes(fig_comp), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    if fig_ahorro is not None:
        doc.add_paragraph("Ahorro sugerido por categoría (Top 10)")
        doc.add_picture(fig_to_png_bytes(fig_ahorro), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================
# HERO
# ==========================================
st.markdown(
    """
    <div class="hero-wrap">
        <div class="pill"><span class="pill-dot"></span> EQ. BTS (Prescriptiva)</div>
        <div class="hero-title">Analítica <span class="accent">Prescriptiva</span></div>
        <div class="hero-sub">
            Este módulo genera un <b>plan de ajuste</b> por categoría para cumplir
            un <b>presupuesto objetivo</b>. Sube tu CSV y descarga el reporte.
        </div>
    </div>
    <div class="sep"></div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# PANEL REQUISITOS
# ==========================================
st.markdown(
    """
    <div class="panel">
        <div style="display:flex; gap:12px;">
            <div style="width:44px;height:44px;border-radius:14px;
                        display:flex;align-items:center;justify-content:center;
                        border:1px solid rgba(34,211,238,.35);
                        background:rgba(2,6,23,.25);color:#22D3EE;">
                <i class="bi bi-lightbulb"></i>
            </div>
            <div>
                <div class="panel-title">Requisitos del archivo</div>
                <small>
                    <b>Columnas requeridas:</b>
                    <code>fecha</code>, <code>monto</code>, <code>categoria</code>
                    &nbsp;&nbsp;|&nbsp;&nbsp;
                    <b>Formato:</b> CSV
                </small>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# CARGA CSV
# ==========================================
file = st.file_uploader("Sube tu archivo CSV", type=["csv"])
if not file:
    st.info("Sube un CSV para comenzar.")
    st.stop()

try:
    df = pd.read_csv(file)
except Exception as e:
    st.error(f"No se pudo leer el CSV. Error: {e}")
    st.stop()

df.columns = [c.strip().lower() for c in df.columns]

required = {"fecha", "monto", "categoria"}
if not required.issubset(df.columns):
    st.error("El CSV debe contener las columnas: fecha, monto y categoria.")
    st.stop()

# Limpieza base
df["categoria"] = df["categoria"].astype(str).str.strip()
df["fecha"] = pd.to_datetime(df["fecha"], errors="coerce")
df["monto"] = pd.to_numeric(df["monto"], errors="coerce")

registros_cargados = len(df)
df = df.dropna(subset=["fecha", "monto", "categoria"]).copy()
df = df[df["categoria"].str.len() > 0].copy()

registros_validos = len(df)
neg_detect = int((df["monto"] < 0).sum())

# ==========================================
# SIDEBAR: Config + escenarios + estado
# ==========================================
st.sidebar.markdown("## Configuración")
remove_negative = st.sidebar.checkbox("Quitar montos negativos", value=True)

st.sidebar.markdown("## Escenarios")
estilo = st.sidebar.radio(
    "Elige un estilo de ajuste",
    ["Balanceado", "Conservador", "Agresivo"],
    index=1
)

st.sidebar.markdown("## Estado del archivo")
neg_removed = 0
if remove_negative:
    neg_removed = int((df["monto"] < 0).sum())
    df = df[df["monto"] >= 0].copy()

st.sidebar.write(f"Registros cargados: {registros_cargados}")
st.sidebar.write(f"Registros válidos: {registros_validos}")
st.sidebar.write(f"Negativos detectados: {neg_detect}")
st.sidebar.write(f"Negativos removidos: {neg_removed}")

if df.empty:
    st.warning("No hay datos válidos después de limpiar el archivo.")
    st.stop()

# ==========================================
# Construir gasto mensual por categoría (desde transacciones)
# ==========================================
df["mes"] = df["fecha"].dt.to_period("M").dt.to_timestamp()

meses = sorted(df["mes"].dropna().unique())
if len(meses) == 0:
    st.warning("No se detectaron meses válidos en 'fecha'.")
    st.stop()

# Sidebar filtro por meses (slider)
st.sidebar.markdown("## Filtro por meses")
min_m, max_m = min(meses), max(meses)

# Slider de meses (mostrando YYYY-MM)
m_labels = [pd.to_datetime(m).strftime("%Y-%m") for m in meses]
i_min_default, i_max_default = 0, len(meses) - 1

r = st.sidebar.slider(
    "Rango de meses",
    min_value=0,
    max_value=len(meses) - 1,
    value=(i_min_default, i_max_default),
)

i1, i2 = r
start_m = meses[i1]
end_m = meses[i2]
meses_label = f"{pd.to_datetime(start_m).strftime('%Y-%m')} a {pd.to_datetime(end_m).strftime('%Y-%m')}"

df_f = df[(df["mes"] >= start_m) & (df["mes"] <= end_m)].copy()
if df_f.empty:
    st.warning("Con ese filtro por meses no quedan datos.")
    st.stop()

# Resumen por categoría (mensual agregado en el rango)
df_resumen = (
    df_f.groupby("categoria", as_index=False)["monto"]
        .sum()
        .rename(columns={"monto": "gasto_mensual"})
        .sort_values("gasto_mensual", ascending=False)
        .reset_index(drop=True)
)

total_actual = float(df_resumen["gasto_mensual"].sum())
n_cats = int(df_resumen["categoria"].nunique())

# ==========================================
# UI: Tabla gasto actual + KPIs
# ==========================================
st.markdown('<div class="section-title">Gasto actual por categoría</div>', unsafe_allow_html=True)
st.dataframe(df_resumen, use_container_width=True)

k1, k2, k3 = st.columns(3)
k1.metric("Total actual", money(total_actual))
k2.metric("Categorías", f"{n_cats}")
k3.metric("Escenario", estilo)

st.divider()

# ==========================================
# Presupuesto objetivo
# ==========================================
st.markdown('<div class="section-title">Presupuesto objetivo</div>', unsafe_allow_html=True)

presupuesto = st.number_input(
    "Define tu presupuesto mensual objetivo ($)",
    min_value=0.0,
    value=float(total_actual),
    step=50.0
)

exceso = max(0.0, float(total_actual) - float(presupuesto))

p1, p2, p3 = st.columns(3)
p1.metric("Presupuesto objetivo", money(presupuesto))
p2.metric("Total actual", money(total_actual))
p3.metric("Ajuste requerido", money(exceso))

st.divider()

# ==========================================
# PLAN PRESCRIPTIVO (escenarios sí influyen)
# ==========================================
st.markdown('<div class="section-title">Plan prescriptivo por categoría</div>', unsafe_allow_html=True)

df_plan = df_resumen.copy()
df_plan["porcentaje_total"] = df_plan["gasto_mensual"] / max(total_actual, 1e-9)

# Límites por escenario (qué tan agresivo permitimos recortar por categoría)
if estilo == "Conservador":
    max_cut_per_cat = 0.12
    focus_top = 3
elif estilo == "Balanceado":
    max_cut_per_cat = 0.20
    focus_top = 5
else:  # Agresivo
    max_cut_per_cat = 0.30
    focus_top = 8

df_plan["reducir_sugerido"] = 0.0
df_plan["objetivo_recomendado"] = df_plan["gasto_mensual"]
df_plan["accion"] = "Mantener"

if exceso <= 0:
    st.success("✅ Ya estás dentro del presupuesto. Mantén límites por categoría y monitorea mes a mes.")
else:
    st.error(f"⚠️ Estás excediendo el presupuesto por: {money(exceso)}")

    # 1) Propuesta base: proporcional por peso
    df_plan["reducir_sugerido"] = df_plan["porcentaje_total"] * exceso

    # 2) Aplicar límite por escenario
    max_allowed = df_plan["gasto_mensual"] * max_cut_per_cat
    df_plan["reducir_sugerido"] = np.minimum(df_plan["reducir_sugerido"], max_allowed)

    # 3) Enfocar el recorte a top categorías si aplica
    df_plan = df_plan.sort_values("reducir_sugerido", ascending=False).reset_index(drop=True)
    mask_focus = np.zeros(len(df_plan), dtype=bool)
    mask_focus[:min(focus_top, len(df_plan))] = True
    df_plan.loc[~mask_focus, "reducir_sugerido"] *= 0.35

    # 4) Recalibrar para alcanzar el exceso (sin pasarse del límite)
    target = exceso
    for _ in range(8):
        current = float(df_plan["reducir_sugerido"].sum())
        if current <= 0:
            break
        ratio = target / current
        df_plan["reducir_sugerido"] *= ratio
        df_plan["reducir_sugerido"] = np.minimum(df_plan["reducir_sugerido"], max_allowed)

    # 5) Objetivos
    df_plan["objetivo_recomendado"] = df_plan["gasto_mensual"] - df_plan["reducir_sugerido"]
    df_plan["accion"] = df_plan["reducir_sugerido"].apply(lambda x: "Reducir" if x >= 1 else "Mantener")

# Prioridad (Top 3 si hay exceso)
df_plan = df_plan.sort_values("reducir_sugerido", ascending=False).reset_index(drop=True)
df_plan["prioridad"] = "Media"
if exceso > 0:
    df_plan.loc[:2, "prioridad"] = "Alta"

# Salidas
df_out = df_plan[[
    "categoria", "gasto_mensual", "porcentaje_total",
    "objetivo_recomendado", "reducir_sugerido", "accion", "prioridad"
]].copy()

df_show = df_out.copy()
df_show["peso"] = (df_show["porcentaje_total"] * 100).round(2).astype(str) + "%"
df_show = df_show.drop(columns=["porcentaje_total"])
df_show["gasto_mensual"] = df_show["gasto_mensual"].apply(money)
df_show["objetivo_recomendado"] = df_show["objetivo_recomendado"].apply(money)
df_show["reducir_sugerido"] = df_show["reducir_sugerido"].apply(money)

st.dataframe(df_show, use_container_width=True)

# ==========================================
# GRÁFICAS (PRO)
# ==========================================
st.markdown('<div class="section-title">Visualización del ajuste</div>', unsafe_allow_html=True)

plot_df = df_out.sort_values("gasto_mensual", ascending=False).head(10).copy()

colG1, colG2 = st.columns(2)

with colG1:
    st.markdown("### Gasto actual vs objetivo recomendado (Top 10)")
    fig, ax = plt.subplots(figsize=(7.4, 4.2))

    x = np.arange(len(plot_df))
    ax.bar(x - 0.2, plot_df["gasto_mensual"], width=0.4, label="Actual")
    ax.bar(x + 0.2, plot_df["objetivo_recomendado"], width=0.4, label="Objetivo")

    ax.set_xticks(x)
    ax.set_xticklabels(plot_df["categoria"], rotation=30, ha="right")
    ax.set_ylabel("Monto")
    ax.set_title("Comparación por categoría")
    style_dark_matplotlib(ax)
    ax.legend()
    plt.tight_layout()
    st.pyplot(fig, use_container_width=True)

with colG2:
    st.markdown("### Ahorro sugerido por categoría (Top 10)")
    fig2, ax2 = plt.subplots(figsize=(7.4, 4.2))
    ax2.bar(plot_df["categoria"], plot_df["reducir_sugerido"])
    ax2.set_ylabel("Ahorro sugerido")
    ax2.set_title("Dónde está el impacto")
    plt.xticks(rotation=30, ha="right")
    style_dark_matplotlib(ax2)
    plt.tight_layout()
    st.pyplot(fig2, use_container_width=True)

# ==========================================
# INTERPRETACIÓN PRO (CARD)
# ==========================================
status_title, status_text, bullets, action_hint, executive = build_prescriptive_explain(
    total_actual=total_actual,
    presupuesto=float(presupuesto),
    exceso=exceso,
    df_plan=df_plan,
    scenario_name=estilo
)

badge_class = "badge-ok" if exceso <= 0 else "badge-warn"
bullets_html = "".join([f"<li>{b}</li>" for b in bullets])

html = dedent(f"""
<div class="interpret-card">
  <div class="interpret-head">
    <div class="interpret-title">Interpretación automática</div>
    <span class="badge {badge_class}">{status_title}</span>
  </div>

  <div class="interpret-exec">{executive}</div>

  <div class="interpret-text">{status_text}</div>

  <ul class="interpret-bullets">
    {bullets_html}
  </ul>

  <div class="interpret-tip">
    <span class="tip-label">Siguiente paso:</span> {action_hint}
  </div>

  <div class="interpret-note">
    Este módulo es prescriptivo porque transforma el diagnóstico (gasto actual)
    en recomendaciones accionables para cumplir un objetivo (presupuesto).
  </div>
</div>
""")

st.markdown(html, unsafe_allow_html=True)

# ==========================================
# DESCARGAS
# ==========================================
st.divider()
st.markdown('<div class="section-title">Descargar reporte</div>', unsafe_allow_html=True)

# ==========================================
# DESCARGA EXCEL (PRESCRIPTIVA + GRÁFICAS)
# ==========================================
from openpyxl.styles import Font, PatternFill
from openpyxl.chart import BarChart, Reference
from io import BytesIO

excel_buffer = BytesIO()

with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
    # ----------------------------------
    # HOJA 1: PLAN PRESCRIPTIVO (NUMÉRICO)
    # ----------------------------------
    plan_excel = df_out.copy()
    plan_excel["porcentaje_total"] = (plan_excel["porcentaje_total"] * 100).round(2)

    plan_excel.to_excel(
        writer,
        sheet_name="Plan_prescriptivo",
        index=False
    )

    # ----------------------------------
    # HOJA 2: RESUMEN ACTUAL POR CATEGORÍA
    # ----------------------------------
    resumen_excel = df_resumen.copy()
    resumen_excel.to_excel(
        writer,
        sheet_name="Gasto_actual",
        index=False
    )

    wb = writer.book

    # ==================================
    # ESTILOS COMUNES
    # ==================================
    header_fill = PatternFill(
        start_color="1F4FD8", end_color="1F4FD8", fill_type="solid"
    )  # azul rey
    header_font = Font(color="FFFFFF", bold=True)

    # ----------------------------------
    # FORMATO PLAN PRESCRIPTIVO
    # ----------------------------------
    ws_p = writer.sheets["Plan_prescriptivo"]
    ws_p.freeze_panes = "A2"

    for cell in ws_p[1]:
        cell.fill = header_fill
        cell.font = header_font

    ws_p.column_dimensions["A"].width = 22  # categoria
    ws_p.column_dimensions["B"].width = 18  # gasto
    ws_p.column_dimensions["C"].width = 18  # %
    ws_p.column_dimensions["D"].width = 22  # objetivo
    ws_p.column_dimensions["E"].width = 18  # reducir
    ws_p.column_dimensions["F"].width = 14  # accion
    ws_p.column_dimensions["G"].width = 14  # prioridad

    for col in ["B", "D", "E"]:
        for cell in ws_p[col][1:]:
            cell.number_format = '"$"#,##0.00'

    for cell in ws_p["C"][1:]:
        cell.number_format = '0.00"%"'

    # ----------------------------------
    # FORMATO RESUMEN
    # ----------------------------------
    ws_r = writer.sheets["Gasto_actual"]
    ws_r.freeze_panes = "A2"

    for cell in ws_r[1]:
        cell.fill = header_fill
        cell.font = header_font

    ws_r.column_dimensions["A"].width = 24
    ws_r.column_dimensions["B"].width = 18

    for cell in ws_r["B"][1:]:
        cell.number_format = '"$"#,##0.00'

    # ==================================
    # GRÁFICA 1: ACTUAL VS OBJETIVO
    # ==================================
    chart_comp = BarChart()
    chart_comp.title = "Gasto actual vs objetivo recomendado"
    chart_comp.y_axis.title = "Monto ($)"
    chart_comp.x_axis.title = "Categoría"

    data = Reference(
        ws_p,
        min_col=2,
        min_row=1,
        max_col=4,
        max_row=ws_p.max_row
    )

    cats = Reference(
        ws_p,
        min_col=1,
        min_row=2,
        max_row=ws_p.max_row
    )

    chart_comp.add_data(data, titles_from_data=True)
    chart_comp.set_categories(cats)
    chart_comp.height = 11
    chart_comp.width = 22

    ws_p.add_chart(chart_comp, "I2")

    # ==================================
    # GRÁFICA 2: AHORRO SUGERIDO
    # ==================================
    chart_save = BarChart()
    chart_save.title = "Ahorro sugerido por categoría"
    chart_save.y_axis.title = "Ahorro ($)"
    chart_save.x_axis.title = "Categoría"

    data = Reference(
        ws_p,
        min_col=5,
        min_row=1,
        max_row=ws_p.max_row
    )

    chart_save.add_data(data, titles_from_data=True)
    chart_save.set_categories(cats)
    chart_save.height = 11
    chart_save.width = 22

    ws_p.add_chart(chart_save, "I20")

excel_buffer.seek(0)

st.download_button(
    "⬇ Descargar Excel prescriptivo (con gráficas)",
    data=excel_buffer.getvalue(),
    file_name="reporte_prescriptivo_plan.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)


# ==========================================
# WORD (GENERAR + DESCARGAR)
# ==========================================
st.markdown(
    """
    <div class="word-card">
      <div class="word-title">Descargar Reporte Word</div>
      <div class="word-sub">
        Genera un archivo .docx con KPIs, plan prescriptivo, gráficas e interpretación completa.
        (Incluye todo lo mostrado en pantalla.)
      </div>
    </div>
    """,
    unsafe_allow_html=True
)

if "word_pres_bytes" not in st.session_state:
    st.session_state.word_pres_bytes = None

colW1, colW2 = st.columns([1, 1.2])

with colW1:
    generar_word = st.button("📄 Generar Word (prescriptiva)")

with colW2:
    if st.session_state.word_pres_bytes:
        st.download_button(
            "⬇ Descargar Reporte Word (.docx)",
            data=st.session_state.word_pres_bytes,
            file_name="reporte_prescriptivo_completo.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.markdown("<div style='height: 52px;'></div>", unsafe_allow_html=True)

if generar_word:
    st.session_state.word_pres_bytes = build_prescriptive_word_report(
        registros_cargados=registros_cargados,
        registros_validos=registros_validos,
        neg_detect=neg_detect,
        neg_removed=neg_removed,
        remove_negative=remove_negative,
        escenario=estilo,
        meses_label=meses_label,
        presupuesto=float(presupuesto),
        total_actual=float(total_actual),
        exceso=float(exceso),
        df_resumen=df_resumen,
        df_show=df_show,
        df_out=df_out,
        status_title=status_title,
        status_text=status_text,
        bullets=bullets,
        action_hint=action_hint,
        executive=executive,
        fig_comp=fig,
        fig_ahorro=fig2
    )

st.caption("© 2025 Portal de Analítica | Módulo Prescriptivo")
